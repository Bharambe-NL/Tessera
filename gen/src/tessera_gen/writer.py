"""Write the corpus to disk. Doc 02 section 8.

The output tree is fixed by the spec::

    synthetic/<seed>/
      entities.json
      facts.jsonl
      corpus/{regulatory,internal,web}/...
      snapshots/T0.json T1.json T2.json T3.json
      questions.jsonl
      boards/<one directory per board>
      ledger.jsonl
      README.md

"Documents are generated as markdown first and rendered to docx, pdf, and html by
deterministic converters." Deterministic is the operative word: doc 02 section 9
requires two builds at one seed to be byte identical, and a PDF that stamps a
creation date is not. Every renderer here has its timestamps pinned.
"""

from __future__ import annotations

import json
from collections.abc import Iterable
from dataclasses import asdict
from pathlib import Path
from typing import Any

from .boards import Board
from .boards import summarise as summarise_boards
from .corpus import Document
from .entities import manifest, synthetic_source_hierarchy
from .facts import Fact
from .facts import summarise as summarise_facts
from .matchers import MATCHERS_VERSION
from .memory import MemoryTruth
from .memory import summarise as summarise_memory
from .mess import Transformation
from .questions import Question, audiences_manifest
from .questions import summarise as summarise_questions
from .rng import GENERATOR_VERSION
from .snapshots import Snapshot

#: Pinned so a rendered document is byte identical between builds.
FIXED_TIMESTAMP = (2026, 1, 1, 0, 0, 0)


def write_json(path: Path, payload: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(
        json.dumps(payload, indent=2, sort_keys=True, ensure_ascii=False) + "\n",
        encoding="utf-8",
    )


def write_jsonl(path: Path, rows: Iterable[Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8", newline="\n") as fh:
        for row in rows:
            fh.write(json.dumps(row, sort_keys=True, ensure_ascii=False) + "\n")


def normalise_zip(path: Path) -> None:
    """Rewrite a zip with fixed member timestamps and a stable member order.

    docx and xlsx are zip archives, and both python-docx and openpyxl stamp each
    member with the wall clock time it was written. Two builds at one seed would
    then differ in bytes while being identical in content, which is exactly the
    drift doc 02 section 9 rules out. Compression level is pinned too, since a
    library default could change under us.
    """
    import re
    import zipfile

    with zipfile.ZipFile(path, "r") as source:
        members = [(info.filename, source.read(info.filename)) for info in source.infolist()]

    # openpyxl overwrites dcterms:modified with the wall clock at save time,
    # whatever the workbook properties said. Put the pinned value back.
    stamp = "{:04d}-{:02d}-{:02d}T{:02d}:{:02d}:{:02d}Z".format(*FIXED_TIMESTAMP)
    pattern = re.compile(
        rb"(<dcterms:(?:created|modified)[^>]*>)[^<]*(</dcterms:(?:created|modified)>)"
    )
    members = [
        (
            name,
            (
                pattern.sub(rb"\g<1>" + stamp.encode() + rb"\g<2>", data)
                if name == "docProps/core.xml"
                else data
            ),
        )
        for name, data in members
    ]

    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=6) as out:
        for name, data in members:
            info = zipfile.ZipInfo(name, date_time=FIXED_TIMESTAMP)
            info.compress_type = zipfile.ZIP_DEFLATED
            info.external_attr = 0o600 << 16
            out.writestr(info, data)


# ---------------------------------------------------------------- renderers --


def render_markdown(doc: Document) -> str:
    return doc.body + "\n"


def render_html(doc: Document) -> str:
    """Web pages are authored as html fragments; everything else is wrapped."""
    if doc.format == "html":
        body = doc.body
    else:
        body = "\n".join(f"<p>{_escape(p.text)}</p>" for p in doc.passages)
    return (
        '<!doctype html>\n<html lang="' + doc.language + '">\n<head>\n'
        '<meta charset="utf-8">\n'
        f"<title>{_escape(doc.title)}</title>\n"
        f"<meta name=\"issuer\" content=\"{_escape(doc.issuer or '')}\">\n"
        f"<meta name=\"published\" content=\"{doc.published_at or ''}\">\n"
        "</head>\n<body>\n" + body + "\n</body>\n</html>\n"
    )


def _escape(s: str) -> str:
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def render_docx(doc: Document, path: Path) -> None:
    from docx import Document as Docx

    d = Docx()
    d.core_properties.title = doc.title
    d.core_properties.author = doc.issuer or "Unknown"
    # python-docx stamps a created/modified time; pin both so builds match.
    import datetime

    stamp = datetime.datetime(*FIXED_TIMESTAMP)
    d.core_properties.created = stamp
    d.core_properties.modified = stamp

    for passage in doc.passages:
        for line in passage.text.split("\n"):
            if line.startswith("## "):
                d.add_heading(line[3:], level=2)
            elif line.startswith("# "):
                d.add_heading(line[2:], level=1)
            elif line.strip():
                d.add_paragraph(line)
    d.save(path)
    normalise_zip(path)


def render_xlsx(doc: Document, path: Path) -> None:
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Exposures"

    for passage in doc.passages:
        for line in passage.text.split("\n"):
            if line.startswith("|") and "---" not in line:
                cells = [c.strip() for c in line.strip("|").split("|")]
                ws.append(cells)
            elif line.strip():
                ws.append([line.strip()])

    # Doc 02 section 5.3's merged headers.
    if ws.max_row >= 1:
        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=2)
    # openpyxl stamps the current time unless told otherwise, which would make
    # two builds at one seed differ. Pin both.
    import datetime

    stamp = datetime.datetime(*FIXED_TIMESTAMP)
    wb.properties.creator = doc.issuer or "Unknown"
    wb.properties.created = stamp
    wb.properties.modified = stamp
    wb.save(path)
    normalise_zip(path)


def render_pdf(doc: Document, path: Path) -> None:
    """A text layer pdf, unless layer 3 marked it as scanned.

    Doc 02 section 5.3: 5 percent of pdfs are rendered to image with no text
    layer, so the Reader's vision path is exercised inside retrieval.
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.pdfgen import canvas as pdf_canvas

    scanned = "scanned_no_text_layer" in doc.transformations

    # invariant pins the document id and the creation date, which reportlab
    # otherwise randomises and stamps. Without it two builds at one seed differ.
    c = pdf_canvas.Canvas(str(path), pagesize=A4, invariant=1)
    c.setTitle(doc.title)
    c.setAuthor(doc.issuer or "Unknown")
    _, height = A4

    y = height - 25 * mm
    for passage in doc.passages:
        for line in passage.text.split("\n"):
            if not line.strip():
                continue
            if y < 25 * mm:
                c.showPage()
                y = height - 25 * mm
            if scanned:
                # No text layer: the glyphs are drawn as paths, so nothing is
                # extractable and the OCR path has to run.
                _draw_as_paths(c, line, 20 * mm, y)
            else:
                c.setFont("Helvetica", 10)
                c.drawString(20 * mm, y, line[:110])
            y -= 6 * mm
    c.save()


def _draw_as_paths(c, text: str, x: float, y: float) -> None:
    """Approximate the shape of a line without any extractable text.

    Enough ink that a vision model has something to read, and no text object at
    all, which is what makes the OCR path necessary rather than optional.
    """
    c.setLineWidth(0.6)
    cursor = x
    for ch in text[:110]:
        if ch == " ":
            cursor += 2.6
            continue
        c.line(cursor, y, cursor + 1.6, y + 4.2)
        c.line(cursor + 1.6, y + 4.2, cursor + 3.2, y)
        cursor += 4.2


RENDERERS = {
    "md": None,  # written directly
    "html": None,
    "txt": None,
    "docx": render_docx,
    "xlsx": render_xlsx,
    "pdf": render_pdf,
}


def write_document(root: Path, doc: Document) -> None:
    path = root / "corpus" / doc.path
    path.parent.mkdir(parents=True, exist_ok=True)

    if doc.format in ("md", "txt"):
        # An empty or corrupt file is a deliberate case, not a bug.
        if "empty" in doc.transformations:
            path.write_text("", encoding="utf-8")
        else:
            path.write_text(render_markdown(doc), encoding="utf-8", newline="\n")
        return

    if doc.format == "html":
        path.write_text(render_html(doc), encoding="utf-8", newline="\n")
        return

    if "corrupt" in doc.transformations:
        # A pdf header followed by nothing a parser can use. Doc 05 section 10
        # `parse_error`: skip the file and record it, never fail the run.
        path.write_bytes(b"%PDF-1.7\n%\xe2\xe3\xcf\xd3\ntruncated on purpose")
        return
    if "password_protected" in doc.transformations:
        path.write_bytes(b"%PDF-1.7\n/Encrypt 1 0 R\ntessera synthetic, no password exists")
        return

    renderer = RENDERERS.get(doc.format)
    if renderer is not None:
        renderer(doc, path)


# ------------------------------------------------------------------- corpus --


def write_corpus(
    root: Path,
    seed: int,
    facts: list[Fact],
    documents: list[Document],
    questions: list[Question],
    boards: list[Board],
    snapshots: list[Snapshot],
    memory_truth: MemoryTruth,
    transformations: list[Transformation],
    dropped_questions: list[dict],
    verification_problems: list[str],
) -> dict:
    """Write the whole tree and return the summary the CLI prints."""
    root.mkdir(parents=True, exist_ok=True)

    write_json(root / "entities.json", manifest())
    write_jsonl(root / "facts.jsonl", (f.to_json() for f in facts))
    write_jsonl(root / "questions.jsonl", (q.to_json() for q in questions))

    for doc in documents:
        write_document(root, doc)
    write_jsonl(root / "corpus" / "documents.jsonl", (d.to_json() for d in documents))

    for snap in snapshots:
        write_json(root / "snapshots" / f"{snap.label}.json", snap.to_json())

    for board in boards:
        write_json(root / "boards" / board.board_id / "board.json", board.to_json())

    # Doc 15 section 5. What the boards retriever should find, what it must not,
    # and the two planted cases the Verifier is scored on. At the root rather
    # than under boards/, because it is not a board and everything that walks
    # that directory expects to find only boards there.
    write_json(root / "memory.json", memory_truth.to_json())

    # The synthetic sibling of the finance pack: the same rules with the
    # synthetic issuers substituted into the source hierarchy (doc 02 section 4).
    write_json(
        root / "finance-eu-synthetic.pack.json",
        {
            "source_hierarchy": synthetic_source_hierarchy(),
            "audiences": audiences_manifest(),
            "note": (
                "The rules come from finance-eu. Only the issuers differ, so a corpus "
                "score is comparable with the shipped pack."
            ),
        },
    )

    summary = {
        "generator_version": GENERATOR_VERSION,
        "matchers_version": MATCHERS_VERSION,
        "seed": seed,
        "corpus_name": f"{GENERATOR_VERSION}-{seed}",
        "facts": summarise_facts(facts),
        "questions": summarise_questions(questions),
        "documents": {
            "total": len(documents),
            "by_kind": _count(d.kind for d in documents),
            "by_format": _count(d.format for d in documents),
            "edge_cases": _count(d.edge_case_id for d in documents if d.edge_case_id),
            "languages": _count(d.language for d in documents),
        },
        "boards": summarise_boards(boards),
        "memory": summarise_memory(memory_truth),
        "snapshots": [s.label for s in snapshots],
        "transformations": _count(t.kind for t in transformations),
        "dropped_questions": len(dropped_questions),
        "verification_problems": len(verification_problems),
    }

    # The ledger: every planting, every transformation, every drop, with the seed.
    write_jsonl(
        root / "ledger.jsonl",
        [
            {"type": "build", **summary},
            *[
                {
                    "type": "planting",
                    "fact_id": f.fact_id,
                    "doc_id": p.doc_id,
                    "passage_id": p.passage_id,
                    "fidelity": p.fidelity,
                }
                for f in facts
                for p in f.planted_in
            ],
            *[{"type": "transformation", **asdict(t)} for t in transformations],
            *[{"type": "dropped_question", **d} for d in dropped_questions],
            *[{"type": "verification_problem", "detail": p} for p in verification_problems],
        ],
    )

    (root / "README.md").write_text(_readme(summary), encoding="utf-8", newline="\n")
    return summary


def _count(values: Iterable[str]) -> dict[str, int]:
    out: dict[str, int] = {}
    for v in values:
        out[v] = out.get(v, 0) + 1
    return dict(sorted(out.items()))


def _readme(summary: dict) -> str:
    """Doc 02 section 8: how this corpus was produced, what it contains, how to
    regenerate."""
    return f"""# Synthetic corpus {summary["corpus_name"]}

Produced by `gen build --seed {summary["seed"]}`. Every name in it is invented;
see `entities.json`. Nothing here describes a real regulator, bank or rule.

## What it contains

- **{summary["facts"]["total"]} facts** across {len(summary["facts"]["by_domain"])} domains,
  {summary["facts"]["numeric_share"]:.0%} of them numbers or dates.
  {summary["facts"]["by_truth"].get("superseded", 0)} are superseded and
  {summary["facts"]["by_truth"].get("false_plant", 0)} are wrong on purpose.
- **{summary["documents"]["total"]} documents**: {summary["documents"]["by_kind"]}.
- **{summary["questions"]["total"]} questions**: {summary["questions"]["by_kind"]},
  {summary["questions"]["advice_bait"]} advice bait and
  {summary["questions"]["empty_corpus"]} about a domain the corpus says nothing about.
- **{summary["boards"]["total"]} boards** at T1, carrying
  {summary["boards"]["cards"]} cards.
- **{len(summary["snapshots"])} snapshots**: {", ".join(summary["snapshots"])}.

## How to regenerate

```
gen build --seed {summary["seed"]} --out synthetic/
gen verify --seed {summary["seed"]}
gen serve --seed {summary["seed"]}
```

The build is deterministic: the same seed produces the same ledger, byte for
byte. `gen verify` re-runs the checks that make the corpus usable, chiefly that
every passage claiming exact fidelity for a fact really does contain its value.

## What it is for

The Verifier's job is to catch unsupported claims, wrong citations, stale sources
and advice language. None of those can be measured on a real corpus without a
human labelling every claim. Here the labels come for free, because the generator
planted the facts, the contradictions and the traps on purpose.

Generator {summary["generator_version"]}, matchers {summary["matchers_version"]}.
Results reference the corpus name so numbers stay comparable across runs.
"""
